#!/usr/bin/env python3
"""
sync_to_drive.py — Mirror every JFrog Learn content page into Google Drive as a
native Google Doc with high fidelity: real headings, native tables, shaded
monospace code, callouts, and EMBEDDED diagram images that look exactly like the
website. Plus an auto-rebuilt index Doc. Designed to run in CI on every push to
main, then have NotebookLM use the Docs as sources.

How it works (the DOCX pipeline):
  1. render_diagrams.py renders each page's CSS/HTML <figure> to a high-DPI PNG
     using headless Chromium, so diagrams are faithful images (the site has no
     image files — diagrams are drawn with CSS/HTML).
  2. build_docx.py converts each page's HTML into a .docx with those PNGs embedded
     and native Word structure (headings, lists, tables, code, callouts).
  3. This script uploads each .docx to Drive with the Google Doc mimeType so Drive
     CONVERTS it into a native Google Doc. Drive's DOCX importer is its
     highest-fidelity converter, so structure + embedded images survive well —
     far better than the old Markdown path, and diagrams finally render.

Why this beats the previous approach: uploading Markdown let Drive's crude
Markdown importer mangle tables/code, and Mermaid code fences never rendered. A
.docx carries embedded image bytes and converts cleanly.

Auth & scope (UNCHANGED from before — no re-consent needed):
  OAuth user credentials with the NON-SENSITIVE 'drive.file' scope. The app acts
  as YOU (so Docs count against your own 15 GB quota) and can only see/manage
  files it creates — so it manages its OWN 'jfrog-learning' folder. Provide a
  single env var GOOGLE_OAUTH_CREDENTIALS containing JSON with client_id,
  client_secret, refresh_token (stored as the GitHub secret of the same name).
  A service account still works only on Workspace/Shared Drives.

  NOTE: We do NOT use the Google Docs API, so the 'documents' scope is NOT needed.
  Embedding images happens inside the .docx before upload — no separate image
  uploads, no extra scope.

Idempotency: the folder is the source of truth. Each run looks up a Doc by exact
name; if found it UPDATES that Doc in place by uploading new .docx media (same
file id, so NotebookLM source links stay valid); if not found it CREATES one.
"""
from __future__ import annotations

import json
import os
import sys
import time
from pathlib import Path

from google.oauth2 import service_account
from google.oauth2.credentials import Credentials as UserCredentials
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaFileUpload

import build_docx as B
import render_diagrams as R

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

REPO_ROOT = Path(__file__).resolve().parent.parent
PAGES_DIR = REPO_ROOT / "pages"
AUTOMATION_DIR = Path(__file__).resolve().parent
DOCX_DIR = AUTOMATION_DIR / "_docx"
DIAGRAMS_DIR = AUTOMATION_DIR / "_diagrams"

FOLDER_NAME = os.environ.get("DRIVE_FOLDER_NAME", "jfrog-learning")
EXPLICIT_FOLDER_ID = os.environ.get("DRIVE_FOLDER_ID", "").strip()

EXCLUDE = {"search"}

GOOGLE_DOC_MIME = "application/vnd.google-apps.document"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
FOLDER_MIME = "application/vnd.google-apps.folder"
SITE_BASE_URL = os.environ.get("SITE_BASE_URL", "https://talorlik.github.io/jfrog-learn")

SCOPES = ["https://www.googleapis.com/auth/drive.file"]

PAGE_ORDER = [
    "fundamentals",
    "replication-federation",
    "build-promotion",
    "release-bundles",
    "rest-api",
    "frogbot",
    "pipelines",
    "access-tokens",
    "kubernetes-helm",
]

DOC_NAME_PREFIX = "JFrog Learn — "
INDEX_DOC_NAME = "JFrog Learn — Index"

TOKEN_URI = "https://oauth2.googleapis.com/token"


# ---------------------------------------------------------------------------
# Auth
# ---------------------------------------------------------------------------

def _oauth_creds_from_json(raw: str):
    info = json.loads(raw)
    missing = [k for k in ("client_id", "client_secret", "refresh_token") if not info.get(k)]
    if missing:
        sys.exit(
            "ERROR: GOOGLE_OAUTH_CREDENTIALS is missing required field(s): "
            + ", ".join(missing)
            + ". It must be JSON with client_id, client_secret, refresh_token."
        )
    return UserCredentials(
        token=None,
        refresh_token=info["refresh_token"],
        client_id=info["client_id"],
        client_secret=info["client_secret"],
        token_uri=info.get("token_uri", TOKEN_URI),
        scopes=SCOPES,
    )


def get_drive():
    """Build a Drive client. Prefer OAuth user creds (works on personal Gmail);
    fall back to a service account (Workspace / Shared Drives only)."""
    oauth = os.environ.get("GOOGLE_OAUTH_CREDENTIALS", "").strip()
    sa_raw = os.environ.get("GOOGLE_SERVICE_ACCOUNT_JSON", "").strip()
    keyfile = os.environ.get("GOOGLE_APPLICATION_CREDENTIALS", "").strip()

    if oauth:
        creds = _oauth_creds_from_json(oauth)
        print("auth: OAuth user credentials")
    elif sa_raw:
        creds = service_account.Credentials.from_service_account_info(
            json.loads(sa_raw), scopes=SCOPES
        )
        print("auth: service account (JSON)")
    elif keyfile and Path(keyfile).exists():
        creds = service_account.Credentials.from_service_account_file(keyfile, scopes=SCOPES)
        print("auth: service account (key file)")
    else:
        sys.exit(
            "ERROR: no credentials. Set GOOGLE_OAUTH_CREDENTIALS (recommended; JSON "
            "with client_id, client_secret, refresh_token), or a service account via "
            "GOOGLE_SERVICE_ACCOUNT_JSON / GOOGLE_APPLICATION_CREDENTIALS."
        )
    return build("drive", "v3", credentials=creds, cache_discovery=False)


# ---------------------------------------------------------------------------
# Drive helpers
# ---------------------------------------------------------------------------

def _q_escape(s: str) -> str:
    return s.replace("\\", "\\\\").replace("'", "\\'")


def find_or_create_folder(drive, name: str) -> str:
    if EXPLICIT_FOLDER_ID:
        try:
            drive.files().get(
                fileId=EXPLICIT_FOLDER_ID, fields="id", supportsAllDrives=True
            ).execute()
            return EXPLICIT_FOLDER_ID
        except HttpError as e:
            status = getattr(getattr(e, "resp", None), "status", None)
            if status in (403, 404):
                print(
                    f"NOTE: DRIVE_FOLDER_ID={EXPLICIT_FOLDER_ID} is not accessible to "
                    f"this app (scope is drive.file, which only sees app-created "
                    f"files). Falling back to an app-owned '{name}' folder."
                )
            else:
                raise
    q = (
        f"name = '{_q_escape(name)}' and mimeType = '{FOLDER_MIME}' "
        f"and trashed = false"
    )
    res = drive.files().list(
        q=q, fields="files(id,name)", spaces="drive",
        supportsAllDrives=True, includeItemsFromAllDrives=True,
    ).execute()
    files = res.get("files", [])
    if files:
        return files[0]["id"]
    meta = {"name": name, "mimeType": FOLDER_MIME}
    folder = drive.files().create(body=meta, fields="id", supportsAllDrives=True).execute()
    print(f"created folder '{name}' -> {folder['id']}")
    return folder["id"]


def find_doc(drive, folder_id: str, name: str):
    q = (
        f"name = '{_q_escape(name)}' and mimeType = '{GOOGLE_DOC_MIME}' "
        f"and '{folder_id}' in parents and trashed = false"
    )
    res = drive.files().list(
        q=q, fields="files(id,name,modifiedTime)", spaces="drive",
        supportsAllDrives=True, includeItemsFromAllDrives=True,
    ).execute()
    files = res.get("files", [])
    return files[0] if files else None


def upsert_doc_from_docx(drive, folder_id: str, name: str, docx_path: Path) -> tuple[str, str]:
    """Create or update a Google Doc by uploading a .docx and asking Drive to
    convert it (mimeType=Google Doc). Returns (file_id, action)."""
    existing = find_doc(drive, folder_id, name)
    for attempt in range(3):
        # MediaFileUpload must be recreated per attempt (stream is consumed).
        media = MediaFileUpload(str(docx_path), mimetype=DOCX_MIME, resumable=False)
        try:
            if existing:
                fid = existing["id"]
                # Update content in place: upload new .docx media; Drive
                # re-converts it into the Google Doc, keeping the same file id.
                drive.files().update(
                    fileId=fid, media_body=media, supportsAllDrives=True,
                    body={"name": name, "mimeType": GOOGLE_DOC_MIME},
                ).execute()
                return fid, "updated"
            else:
                meta = {
                    "name": name,
                    "mimeType": GOOGLE_DOC_MIME,  # convert on upload
                    "parents": [folder_id],
                }
                created = drive.files().create(
                    body=meta, media_body=media, fields="id",
                    supportsAllDrives=True,
                ).execute()
                return created["id"], "created"
        except HttpError as e:
            status = getattr(e, "status_code", None) or getattr(e.resp, "status", None)
            if status in (429, 500, 502, 503) and attempt < 2:
                time.sleep(2 ** attempt)
                continue
            raise
    raise RuntimeError(f"failed to upsert '{name}'")


# ---------------------------------------------------------------------------
# Index Doc (built as a .docx too, for consistent styling + working links)
# ---------------------------------------------------------------------------

def build_index_docx(entries: list[dict], out_path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("JFrog Learn — Index", level=0)
    p = doc.add_paragraph()
    p.add_run(
        "Auto-generated index of all JFrog Learn pages, kept in sync with the "
        "website on every push. Each entry links to its Google Doc (use these as "
        "NotebookLM sources) and to the live page."
    ).italic = True

    doc.add_heading("Pages", level=1)
    for i, e in enumerate(entries, 1):
        doc_link = f"https://docs.google.com/document/d/{e['file_id']}/edit"
        para = doc.add_paragraph(style="List Number")
        r = para.add_run(e["title"])
        r.bold = True
        para.add_run("  —  ")
        B._add_hyperlink(para, doc_link, "Google Doc")
        para.add_run("  ·  ")
        B._add_hyperlink(para, e["site_url"], "live page")

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.add_run(
        f"Total pages: {len(entries)}. Folder: '{FOLDER_NAME}'."
    ).italic = True

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# ---------------------------------------------------------------------------
# Page discovery + build
# ---------------------------------------------------------------------------

def list_pages() -> list[str]:
    slugs = []
    for p in sorted(PAGES_DIR.glob("*.html")):
        slug = p.stem
        if slug in EXCLUDE:
            continue
        slugs.append(slug)
    ordered = [s for s in PAGE_ORDER if s in slugs]
    extras = [s for s in slugs if s not in PAGE_ORDER]
    return ordered + extras


def doc_name_for(title: str) -> str:
    return f"{DOC_NAME_PREFIX}{title}"


def render_all_diagrams(slugs: list[str]) -> dict:
    """Render diagrams for all pages and return the manifest dict."""
    import asyncio
    print("Rendering diagrams with headless Chromium ...")
    DIAGRAMS_DIR.mkdir(parents=True, exist_ok=True)
    manifest = asyncio.run(R.main_async(slugs))
    (DIAGRAMS_DIR / "manifest.json").write_text(
        json.dumps(manifest, indent=2), encoding="utf-8"
    )
    total = sum(len(v) for v in manifest.values())
    print(f"Rendered {total} diagram(s).")
    return manifest


def build_all_docx(slugs: list[str], manifest: dict) -> dict[str, dict]:
    """Build a .docx for each page. Returns {slug: {title, docx_path, site_url}}."""
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    out = {}
    for slug in slugs:
        html = (PAGES_DIR / f"{slug}.html").read_text(encoding="utf-8")
        title = B.page_title(html)
        site_url = f"{SITE_BASE_URL}/pages/{slug}.html"
        docx_path = DOCX_DIR / f"{slug}.docx"
        B.build_page_docx(html, manifest.get(slug, []), docx_path, title, site_url)
        out[slug] = {"title": title, "docx_path": docx_path, "site_url": site_url}
        print(f"  built {slug}.docx (title={title!r}, "
              f"diagrams={len(manifest.get(slug, []))})")
    return out


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    dry = "--dry-run" in sys.argv

    slugs = list_pages()
    print(f"pages to sync ({len(slugs)}): {', '.join(slugs)}")

    manifest = render_all_diagrams(slugs)
    built = build_all_docx(slugs, manifest)

    if dry:
        print("\nDRY RUN: built .docx files in automation/_docx/ "
              "and diagram PNGs in automation/_diagrams/. No Drive writes performed.")
        for slug in slugs:
            print(f"  [dry] {slug}: {built[slug]['docx_path']}")
        return

    drive = get_drive()
    folder_id = find_or_create_folder(drive, FOLDER_NAME)
    print(f"folder '{FOLDER_NAME}' -> {folder_id}")

    entries = []
    for slug in slugs:
        info = built[slug]
        name = doc_name_for(info["title"])
        fid, action = upsert_doc_from_docx(drive, folder_id, name, info["docx_path"])
        print(f"  {action}: {name} -> {fid}")
        entries.append({
            "title": info["title"],
            "file_id": fid,
            "site_url": info["site_url"],
        })

    # Rebuild the index doc last (so it has all fresh file ids)
    idx_path = DOCX_DIR / "_index.docx"
    build_index_docx(entries, idx_path)
    idx_id, idx_action = upsert_doc_from_docx(drive, folder_id, INDEX_DOC_NAME, idx_path)
    print(f"  {idx_action}: {INDEX_DOC_NAME} -> {idx_id}")

    print(f"\nDone. {len(entries)} page Docs + 1 index in folder '{FOLDER_NAME}'.")
    print(f"Folder: https://drive.google.com/drive/folders/{folder_id}")


if __name__ == "__main__":
    main()
