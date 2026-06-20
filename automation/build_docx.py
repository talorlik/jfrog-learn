#!/usr/bin/env python3
"""
build_docx.py — Convert a JFrog Learn page (HTML) into a structured .docx with
native headings, lists, tables, shaded monospace code blocks, callouts, and
EMBEDDED diagram images (rendered separately by render_diagrams.py).

Why DOCX (not Markdown, not the Docs API): Google Drive's DOCX importer is its
highest-fidelity converter, so uploading a .docx and letting Drive convert it to
a Google Doc preserves structure far better than the old Markdown path — and a
.docx embeds image bytes directly, so diagrams travel inside the file. This needs
only the drive.file scope (no documents scope, no re-consent).

This module walks the page DOM in document order so diagram images land in the
right place relative to the surrounding text. It reuses the inline-formatting and
table/list/callout understanding from html_to_markdown for text fidelity, but
emits native DOCX constructs instead of Markdown.
"""
from __future__ import annotations

import re
from pathlib import Path

from bs4 import BeautifulSoup, Comment, NavigableString, Tag
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

# ---------------------------------------------------------------------------
# Styling constants
# ---------------------------------------------------------------------------

CODE_FONT = "Consolas"
CODE_BG = "F0F0F2"          # light gray shading behind code blocks
CODE_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
LINK_COLOR = RGBColor(0x1A, 0x57, 0xB7)
CALLOUT_BG = {
    "callout-tip": "EAF6EC",
    "callout-warn": "FCEFE3",
    "callout-note": "EAF0FA",
    "callout-q": "F1ECF8",
}
CALLOUT_LABEL = {
    "callout-tip": "Tip",
    "callout-warn": "Warning",
    "callout-note": "Note",
    "callout-q": "Question",
}
# Max image width in inches for the Doc body (keeps wide diagrams readable).
MAX_IMG_WIDTH_IN = 6.3

SKIP_CLASSES = {"nav-link", "sidebar", "search", "tab-bar", "kicker",
                "theme-toggle", "page-nav", "breadcrumb", "prevnext"}


# ---------------------------------------------------------------------------
# Low-level DOCX helpers
# ---------------------------------------------------------------------------

def _shade(element, fill_hex: str):
    """Apply background shading to a paragraph or table cell <w:...Pr>."""
    pr = element.get_or_add_pPr() if element.tag.endswith("}p") else element
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pr.append(shd)


def _para_shade(paragraph, fill_hex: str):
    _shade(paragraph._p, fill_hex)


def _cell_shade(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "C8C8CE")
        borders.append(el)
    tcPr.append(borders)


def _add_hyperlink(paragraph, url: str, text: str):
    """Insert a clickable hyperlink run into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1A57B7")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ---------------------------------------------------------------------------
# Inline run emission (bold/italic/code/links) onto a paragraph
# ---------------------------------------------------------------------------

def _clean(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def emit_inline(paragraph, node, *, bold=False, italic=False, code=False):
    """Walk inline children and add runs with formatting onto `paragraph`."""
    for child in node.children:
        if isinstance(child, Comment):
            continue
        if isinstance(child, NavigableString):
            txt = str(child)
            if txt:
                run = paragraph.add_run(txt)
                run.bold = bold or None
                run.italic = italic or None
                if code:
                    run.font.name = CODE_FONT
                    run.font.color.rgb = CODE_TEXT
            continue
        if not isinstance(child, Tag):
            continue
        name = child.name.lower()
        if name in ("strong", "b"):
            emit_inline(paragraph, child, bold=True, italic=italic, code=code)
        elif name in ("em", "i"):
            emit_inline(paragraph, child, bold=bold, italic=True, code=code)
        elif name in ("code", "kbd"):
            emit_inline(paragraph, child, bold=bold, italic=italic, code=True)
        elif name == "a":
            href = (child.get("href") or "").strip()
            label = _clean(child.get_text(" ")) or href
            if href and not href.startswith("#"):
                _add_hyperlink(paragraph, href, label)
            else:
                run = paragraph.add_run(label)
                run.bold = bold or None
        elif name == "br":
            paragraph.add_run().add_break()
        else:  # span, small, sup, sub, etc. -> keep text
            emit_inline(paragraph, child, bold=bold, italic=italic, code=code)


def _has_text(node) -> bool:
    return bool(_clean(node.get_text(" ")))


# ---------------------------------------------------------------------------
# Block emitters
# ---------------------------------------------------------------------------

HEADING_LEVEL = {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}


class DocBuilder:
    def __init__(self, doc: Document, diagrams: list[dict]):
        self.doc = doc
        # diagrams: list of {index, abs_path, caption, width, height} for THIS page,
        # consumed in document order as <figure> elements are encountered.
        self.diagrams = diagrams
        self._fig_cursor = 0

    # -- code ---------------------------------------------------------------
    def add_code_block(self, text: str):
        text = text.replace("\xa0", " ").rstrip("\n")
        for i, line in enumerate(text.split("\n")):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            _para_shade(p, CODE_BG)
            run = p.add_run(line if line else " ")
            run.font.name = CODE_FONT
            run.font.size = Pt(9.5)
            run.font.color.rgb = CODE_TEXT
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # -- table --------------------------------------------------------------
    def add_table(self, table_el: Tag):
        head = table_el.find("thead")
        body = table_el.find("tbody") or table_el
        rows = body.find_all("tr")
        headers = []
        if head:
            headers = head.find_all("th")
        if not headers and rows and rows[0].find("th"):
            headers = rows[0].find_all("th")
            rows = rows[1:]
        ncol = max(
            [len(headers)] + [len(r.find_all(["td", "th"])) for r in rows] or [1]
        )
        if ncol == 0:
            return
        t = self.doc.add_table(rows=0, cols=ncol)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.style = "Table Grid"
        if headers:
            hr = t.add_row().cells
            for i in range(ncol):
                cell = hr[i]
                _cell_shade(cell, "E8E8EC")
                cell.paragraphs[0].text = ""
                if i < len(headers):
                    emit_inline(cell.paragraphs[0], headers[i])
                for r in cell.paragraphs[0].runs:
                    r.bold = True
        for tr in rows:
            cells = tr.find_all(["td", "th"])
            if not cells:
                continue
            row = t.add_row().cells
            for i in range(ncol):
                if i < len(cells):
                    row[i].paragraphs[0].text = ""
                    emit_inline(row[i].paragraphs[0], cells[i])
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # -- lists --------------------------------------------------------------
    def add_list(self, lst: Tag, depth: int = 0):
        ordered = lst.name.lower() == "ol"
        style = "List Number" if ordered else "List Bullet"
        if depth > 0:
            style = f"{style} {min(depth+1,3)}"
        for li in lst.find_all("li", recursive=False):
            for n in li.find_all(class_=re.compile(r"ls-n|step-n|card-num")):
                n.decompose()
            nested = [s.extract() for s in li.find_all(["ul", "ol"], recursive=False)]
            try:
                p = self.doc.add_paragraph(style=style)
            except KeyError:
                p = self.doc.add_paragraph(style="List Bullet")
            emit_inline(p, li)
            for sub in nested:
                self.add_list(sub, depth + 1)

    # -- callout ------------------------------------------------------------
    def add_callout(self, div: Tag):
        cls = div.get("class", [])
        kind = next((c for c in cls if c in CALLOUT_BG), None)
        bg = CALLOUT_BG.get(kind, "F2F2F4")
        label = CALLOUT_LABEL.get(kind, "")
        title_el = div.find(class_="callout-title")
        title = _clean(title_el.get_text(" ")) if title_el else ""
        if title_el:
            title_el.extract()
        head_txt = " — ".join(x for x in [label, title] if x)
        if head_txt:
            p = self.doc.add_paragraph()
            _para_shade(p, bg)
            r = p.add_run(head_txt)
            r.bold = True
        # body paragraphs
        for child in div.find_all(["p"], recursive=False) or [div]:
            if not _has_text(child):
                continue
            p = self.doc.add_paragraph()
            _para_shade(p, bg)
            emit_inline(p, child)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # -- glossary entry (dt/dd) --------------------------------------------
    def add_gloss(self, div: Tag):
        dt = div.find("dt")
        dd = div.find("dd")
        if not dt and not dd:
            return
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        if dt:
            r = p.add_run(_clean(dt.get_text(" ")))
            r.bold = True
        if dt and dd:
            p.add_run(" \u2014 ")
        if dd:
            emit_inline(p, dd)

    # -- analogy box (shaded, like a callout) ------------------------------
    def add_analogy(self, div: Tag):
        for ic in div.find_all(class_=re.compile(r"analogy-icon")):
            ic.decompose()
        title_el = div.find(class_="analogy-title")
        title = _clean(title_el.get_text(" ")) if title_el else ""
        if title_el:
            title_el.extract()
        bg = "FFF6E6"
        if title:
            p = self.doc.add_paragraph()
            _para_shade(p, bg)
            p.add_run(title).bold = True
        for child in div.find_all(["p"], recursive=True):
            if not _has_text(child):
                continue
            p = self.doc.add_paragraph()
            _para_shade(p, bg)
            emit_inline(p, child)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # -- numbered concept card ---------------------------------------------
    def add_card(self, div: Tag):
        num_el = div.find(class_=re.compile(r"card-num"))
        num = _clean(num_el.get_text(" ")) if num_el else ""
        if num_el:
            num_el.decompose()
        heading = div.find(["h2", "h3", "h4"])
        if heading:
            htxt = _clean(heading.get_text(" "))
            lvl = HEADING_LEVEL.get(heading.name.lower(), 3)
            label = f"{num}. {htxt}" if num else htxt
            self.doc.add_heading(label, level=min(lvl, 4))
            heading.decompose()
        # remaining body
        for child in div.children:
            if isinstance(child, Tag):
                self.emit_block(child)

    # -- chip / pill grid (render as one compact line) ---------------------
    def add_chips(self, div: Tag):
        chips = [
            _clean(c.get_text(" "))
            for c in div.find_all(class_=re.compile(r"\bchip\b|\bpill\b"))
        ]
        chips = [c for c in chips if c]
        if not chips:
            return
        p = self.doc.add_paragraph()
        p.add_run("  \u00b7  ".join(chips))

    # -- diagram (embed pre-rendered PNG) ----------------------------------
    def add_diagram(self, fig: Tag):
        # Find the matching rendered image by document order.
        img = None
        if self._fig_cursor < len(self.diagrams):
            img = self.diagrams[self._fig_cursor]
        self._fig_cursor += 1
        cap = ""
        cap_el = fig.find("figcaption")
        if cap_el:
            cap = _clean(cap_el.get_text(" "))
        if img and Path(img["abs_path"]).exists():
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            # scale to max width preserving aspect
            w = img.get("width") or 1000
            h = img.get("height") or 600
            width_in = min(MAX_IMG_WIDTH_IN, w / 96.0)
            run.add_picture(img["abs_path"], width=Inches(width_in))
            if cap:
                cp = self.doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cp.add_run(cap)
                cr.italic = True
                cr.font.size = Pt(9)
        else:
            # No image available -> readable text fallback so nothing is lost.
            if cap:
                p = self.doc.add_paragraph()
                p.add_run(f"[Diagram] {cap}").italic = True

    # -- generic walk -------------------------------------------------------
    def walk(self, parent: Tag):
        for child in parent.children:
            if isinstance(child, Comment):
                continue
            if isinstance(child, NavigableString):
                txt = _clean(str(child))
                if txt:
                    self.doc.add_paragraph().add_run(txt)
                continue
            if not isinstance(child, Tag):
                continue
            self.emit_block(child)

    def emit_block(self, el: Tag):
        name = el.name.lower()
        cls = set(el.get("class", []))
        if cls & SKIP_CLASSES:
            return

        if name in HEADING_LEVEL:
            txt = _clean(el.get_text(" "))
            if txt:
                lvl = HEADING_LEVEL[name]
                self.doc.add_heading(txt, level=min(lvl, 4))
            return

        # chip/pill grids -> compact dotted line (before generic div recursion)
        if "chips" in cls or "pills" in cls:
            self.add_chips(el)
            return

        # numbered concept card
        if "card" in cls and "cards" not in cls:
            self.add_card(el)
            return

        # glossary entry
        if "gloss" in cls:
            self.add_gloss(el)
            return

        # analogy box
        if "analogy" in cls:
            self.add_analogy(el)
            return

        # diagram figures
        if name == "figure" or (cls & {"diagram"}) or el.find(
            class_=re.compile(r"flow-row|diag-grid|bm-box|onion")
        ):
            if name == "figure" or (cls & {"diagram", "flow", "bigmodel"}):
                self.add_diagram(el)
                return

        if name == "table":
            self.add_table(el)
            return
        if "datatable" in cls or "table-wrap" in cls:
            tbl = el.find("table")
            if tbl:
                self.add_table(tbl)
                return

        if name in ("ul", "ol"):
            self.add_list(el)
            return

        if name == "pre" or "codeblock" in cls:
            code = el.find("code") or el
            self.add_code_block(code.get_text())
            return

        if "callout" in cls:
            self.add_callout(el)
            return

        if name == "blockquote":
            p = self.doc.add_paragraph(style="Intense Quote")
            emit_inline(p, el)
            return

        if name == "p" or "prose" in cls or "subhead" in cls or "panel-intro" in cls:
            if not _has_text(el):
                return
            if "subhead" in cls:
                p = self.doc.add_paragraph()
                emit_inline(p, el)
                for r in p.runs:
                    r.bold = True
            else:
                p = self.doc.add_paragraph()
                emit_inline(p, el)
            return

        # generic container -> recurse
        if name in ("div", "section", "article", "main", "aside", "details", "header", "footer"):
            self.walk(el)
            return

        # fallback inline
        if _has_text(el):
            p = self.doc.add_paragraph()
            emit_inline(p, el)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def page_title(html: str) -> str:
    soup = BeautifulSoup(html, "lxml")
    h1 = soup.find("h1")
    if h1 and _clean(h1.get_text(" ")):
        return _clean(h1.get_text(" "))
    if soup.title:
        return _clean(soup.title.get_text()).replace(" — JFrog Learn", "")
    return "Untitled"


def build_page_docx(html: str, diagrams: list[dict], out_path: Path,
                    title: str, source_url: str) -> Path:
    """Build a .docx for one page. `diagrams` are this page's rendered figures in
    document order (from render_diagrams.py manifest)."""
    soup = BeautifulSoup(html, "lxml")
    content = soup.find("main") or soup.find("article") or soup.body

    # strip chrome
    for sel in ["nav", "aside", "header", "footer", "script", "style", "form"]:
        for t in content.find_all(sel):
            t.decompose()
    for t in content.find_all(class_=re.compile(
        r"sidebar|search|tab-bar|theme-toggle|page-nav|breadcrumb|pagenav|prevnext")):
        t.decompose()

    doc = Document()
    # Title + source line
    doc.add_heading(title, level=0)
    src_p = doc.add_paragraph()
    src_p.add_run("Source: ").italic = True
    _add_hyperlink(src_p, source_url, title)

    builder = DocBuilder(doc, diagrams)
    builder.walk(content)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


if __name__ == "__main__":
    import json
    import sys
    slug = sys.argv[1] if len(sys.argv) > 1 else "fundamentals"
    repo = Path(__file__).resolve().parent.parent
    html = (repo / "pages" / f"{slug}.html").read_text(encoding="utf-8")
    manifest_path = Path(__file__).resolve().parent / "_diagrams" / "manifest.json"
    manifest = json.loads(manifest_path.read_text()) if manifest_path.exists() else {}
    title = page_title(html)
    out = Path(__file__).resolve().parent / "_docx" / f"{slug}.docx"
    build_page_docx(
        html, manifest.get(slug, []), out, title,
        f"https://talorlik.github.io/jfrog-learn/pages/{slug}.html",
    )
    print(f"wrote {out} (title={title!r}, diagrams={len(manifest.get(slug, []))})")
